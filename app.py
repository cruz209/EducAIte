from flask import Flask, render_template, request, session, jsonify
from flask_sqlalchemy import SQLAlchemy
from transformers import AutoTokenizer, AutoModelForQuestionAnswering
from sentence_transformers import SentenceTransformer, util
from PyPDF2 import PdfReader
import docx
import os
import uuid
import torch

# Initialize Flask app
app = Flask(__name__)
app.secret_key = 'your_secret_key'
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# SQLite configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///documents.db'
db = SQLAlchemy(app)

# Initialize sentence transformer for semantic search
semantic_model = SentenceTransformer('all-MiniLM-L6-v2')

# Initialize BERT for question answering (new model)
bert_model = AutoModelForQuestionAnswering.from_pretrained('bert-large-uncased-whole-word-masking-finetuned-squad')
bert_tokenizer = AutoTokenizer.from_pretrained('bert-large-uncased-whole-word-masking-finetuned-squad')

# Allowed file types for upload
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# Document model for SQLite
class Document(db.Model):
    id = db.Column(db.String(36), primary_key=True)
    chunks = db.Column(db.PickleType, nullable=False)  # Store document chunks with their embeddings

# Create the database and tables if they do not exist
if not os.path.exists('documents.db'):
    with app.app_context():
        db.create_all()

# Helper function to check allowed file types
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to extract text from PDF
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Helper function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Helper function to extract text from TXT
def extract_text_from_txt(file_path):
    with open(file_path, 'r') as file:
        return file.read()

# Helper function to split text into chunks
def split_into_chunks(text, chunk_size=100):
    words = text.split()
    chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
    return chunks

# Route for file upload and processing
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file part", 400
    file = request.files['file']

    if file.filename == '':
        return "No selected file", 400

    if file and allowed_file(file.filename):
        filename = file.filename
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Extract text based on the file type
        if filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            extracted_text = extract_text_from_docx(file_path)
        elif filename.endswith('.txt'):
            extracted_text = extract_text_from_txt(file_path)

        # Split text into chunks
        text_chunks = split_into_chunks(extracted_text)

        # Create embeddings for each chunk
        chunk_embeddings = []
        for chunk in text_chunks:
            embedding = semantic_model.encode(chunk, convert_to_tensor=True)
            chunk_embeddings.append((chunk, embedding))

        # Generate a unique ID for the document
        doc_id = str(uuid.uuid4())

        # Store the chunks and their embeddings in the database
        new_doc = Document(id=doc_id, chunks=chunk_embeddings)
        db.session.add(new_doc)
        db.session.commit()

        # Store the document ID in the session
        session['document_id'] = doc_id

        return "File uploaded and processed successfully", 200

# Route to handle question asking
@app.route('/ask_question', methods=['POST'])
def ask_question():
    question = request.form['question']

    # Check if document ID exists in session
    if 'document_id' not in session:
        return jsonify({"error": "No document uploaded yet"}), 400

    # Get the document ID from the session
    doc_id = session['document_id']

    # Retrieve the document chunks from the database
    document = Document.query.get(doc_id)
    if document is None:
        return jsonify({"error": "Document not found"}), 404

    # Find the top 3 most relevant chunks using semantic search
    query_embedding = semantic_model.encode(question, convert_to_tensor=True)
    scores_chunks = []

    for chunk, embedding in document.chunks:
        score = util.pytorch_cos_sim(query_embedding, embedding).item()
        scores_chunks.append((score, chunk))

    # Sort by score and take the top 3 chunks
    top_chunks = sorted(scores_chunks, key=lambda x: x[0], reverse=True)[:3]
    combined_context = " ".join([chunk for _, chunk in top_chunks])

    # Use BERT model to find the answer in the combined context
    inputs = bert_tokenizer(question, combined_context, return_tensors='pt', truncation=True)
    with torch.no_grad():
        outputs = bert_model(**inputs)

    # Extract the most likely start and end positions of the answer
    answer_start = torch.argmax(outputs.start_logits)
    answer_end = torch.argmax(outputs.end_logits) + 1

    # Convert token IDs back to string (the answer)
    answer = bert_tokenizer.convert_tokens_to_string(bert_tokenizer.convert_ids_to_tokens(inputs['input_ids'][0][answer_start:answer_end]))

    return jsonify({"answer": answer, "quote": combined_context})

# Route to render the main page
@app.route('/')
def index():
    return render_template('index.html')

# Run the Flask app
if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)
